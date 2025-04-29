from OpenOrchestrator.orchestrator_connection.connection import OrchestratorConnection
import uuid
import json
import requests
from datetime import datetime
from SendSMTPMail import send_email 
import pyodbc  
import pandas as pd
import os
from docx import Document
import getpass
import locale
import traceback
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains

def invoke_SendDigitalPost(Arguments_SendDigitalPost,orchestrator_connection: OrchestratorConnection): 
    KMDNovaLoginRykkerRobot = orchestrator_connection.get_credential("KMDNovaRobotLogin")
    RobotUserName = KMDNovaLoginRykkerRobot.username
    RobotPassword = KMDNovaLoginRykkerRobot.password

    danish_months = [
        "januar", "februar", "marts", "april", "maj", "juni",
        "juli", "august", "september", "oktober", "november", "december"
    ]
    Afgørelsesdato = Arguments_SendDigitalPost.get("in_Afgørelsesdato")
    date_obj = datetime.strptime(Afgørelsesdato, "%Y-%m-%dT%H:%M:%S")
    # Format manually for Danish output
    day = date_obj.day
    month = danish_months[date_obj.month - 1]
    year = date_obj.year
    DecisionDate = f"{day}. {month} {year}"    
    
    BeskrivelseTilEjer = Arguments_SendDigitalPost.get("in_BeskrivelseTilEjer")
    Sagsnummer = Arguments_SendDigitalPost.get("in_Sagsnummer")
    Dato = Arguments_SendDigitalPost.get("in_Dato")
    KMDNovaURL = Arguments_SendDigitalPost.get("in_NovaAPIURL")
    RykkerNummer = Arguments_SendDigitalPost.get("in_RykkerNummer")
    Token = Arguments_SendDigitalPost.get("in_Token")
    fullName = Arguments_SendDigitalPost.get("in_fullName")
    racfId = Arguments_SendDigitalPost.get("in_racfId")
    
    NewDate = datetime.now().strftime("%d-%m-%Y")

    def update_word_template(
        EjerNavn,
        EjerAdresse,
        CaseTitle,
        CaseAdress,
        NewDate,
        CaseNumber,
        DecisionDate,
        input_filename
    ):
        replacements = {
            "<<sagEjer1Navn>>": EjerNavn,
            "<<sagEjer1Adresse>>": EjerAdresse,
            "<<sagsNavn>>": CaseTitle,
            "<<sagsAdresse>>": CaseAdress,
            "<<DatoPlaceholder>>": NewDate,
            "«Sagsnummer»": CaseNumber,
            "<<dato>>": DecisionDate
        }

        def replace_in_paragraph(paragraph, replacements):
            full_text = ''.join(run.text for run in paragraph.runs)
            replaced = False
            for key, val in replacements.items():
                if key in full_text:
                    full_text = full_text.replace(key, val)
                    replaced = True
            if replaced:
                for run in paragraph.runs:
                    run.text = ''
                paragraph.runs[0].text = full_text

        def replace_in_container(container, replacements):
            for paragraph in container.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_container(cell, replacements)

        # Load the document
        doc = Document(input_filename)

        # Replace in body
        replace_in_container(doc, replacements)

        # Replace in headers and footers
        for section in doc.sections:
            for hdr in [section.header, section.first_page_header, section.even_page_header]:
                replace_in_container(hdr, replacements)
            for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
                replace_in_container(ftr, replacements)

        # Save to Downloads
        username = getpass.getuser()
        save_path = os.path.join("C:\\Users", username, "Downloads", input_filename)
        doc.save(save_path)
        print(f"✅ Document saved to: {save_path}")
            
    
    DocumentSendt = False
    AarhusKommuneCVR = "55133018"
    IsCvrAarhusKommune = False

    # Henter yderligere sags information: 

    TransactionID = str(uuid.uuid4())

    Payload = {
        "common": {
            "transactionId": TransactionID
        },
        "paging": {
            "startRow": 1,
            "numberOfRows": 500
        },
        "caseAttributes": {
            "userFriendlyCaseNumber": Sagsnummer
        },
        "caseGetOutput": {
            "caseAttributes": {
                "title": True
            },
            "caseParty": {
                "index": True,
                "identificationType": True,
                "identification": True,
                "partyRole": True,
                "partyRoleName": True,
                "participantRole": True,
                "name": True,
                "participantContactInformation": True,
                "participantRemark": True
            },
            "buildingCase": {
                "propertyInformation": {
                    "caseAddress": True
                }
            }
        }
    }

    # Define headers
    headers = {
        "Authorization": f"Bearer {Token}",
        "Content-Type": "application/json"
    }

    # Define the API endpoint
    url = f"{KMDNovaURL}/Case/GetList?api-version=2.0-Case"
    # Make the HTTP request
    try:
        response = requests.put(url, headers=headers, json=Payload)
        response.raise_for_status()  # Raise an error for non-2xx responses
        data = response.json()
        print("Success:", response.status_code)
    except Exception as e:
        raise Exception("Kunne ikke hente sagsinfo:", str(e))


    # Extract required variables
    for case in data['cases']:
        CaseUuid = case['common']['uuid']
        CaseTitle = case['caseAttributes']['title']
        CaseAddress = case['buildingCase']['propertyInformation']['caseAddress']
        
        identificationType = None
        identification = None

        # Iterate through case parties
        for party in case['caseParties']:
            if "PRI" in party.get('partyRole', '') and "Primær" in party.get('participantRole', ''):
                identificationType = party.get('identificationType')
                identification = party.get('identification')

        # Print the extracted values
        print(f"CaseUuid: {CaseUuid}")
        print(f"CaseTitle: {CaseTitle}")
        print(f"CaseAddress: {CaseAddress}")
        
        if identificationType and identification:
            print(f"IdentificationType: {identificationType}")
            print(f"Identification: {identification}")


    # Get Ejer Oplysninger
    conn_str = "DRIVER={ODBC Driver 17 for SQL Server};SERVER=srvsql29;DATABASE=LOIS;Trusted_Connection=yes"
    conn = pyodbc.connect(conn_str)
    cursor = conn.cursor()

    # Read SQL query from file
    with open('BFENummer.sql', 'r', encoding='utf-8') as file:
        sql_query = file.read()

    sql_query = sql_query.replace("'Ejendomsnummer'", f"'{identification}'")

    # Execute query and fetch results
    cursor.execute(sql_query)
    columns = [column[0] for column in cursor.description]  # Get column names
    rows = cursor.fetchall()

    # Convert to a pandas DataFrame (acting as a DataTable)
    df = pd.DataFrame.from_records(rows, columns=columns)

    # Close database connection
    cursor.close()
    conn.close()

    # Check if DataFrame is empty
    if df.empty:
        raise Exception("Det er hverken et CPR eller CVR-nummer")  # Raise exception if no rows

    print("Der findes en sag med dette BFE-nummer")  # If rows exist

    # Initialize output variables
    CVR = None
    EjerAdresse = None
    EjerNavn = None
    BoolCVR = False
    BoolCPR = False
    BoolBeskyttet = False
    

    # Iterate through DataFrame rows
    for index, row in df.iterrows():
        if "Virksomhed" in str(row["EjerType"]): #Check for EjerType
            print("Det er en virksomhed")
            if "Hovedejer" in str(row["EjerStatusKode_T"]):  # Check if Hovedejer
                CVR = str(row["EjendeVirksomhedCVRnr"])
                EjerAdresse = str(row["EjersAdresseDanmark"])
                EjerNavn = str(row["NavnJusteret"])
                BoolCVR = True

                # Check for Beskyttelse (NULL check)
                if pd.isnull(row["Beskyttelse"]):
                    print("Ejeren har ingen adressebeskyttelse")
                    BoolBeskyttet = False
                else:
                    try:
                        if int(row["Beskyttelse"]) == 1:
                            print("Ejeren har adressebeskyttelse")
                            BoolBeskyttet = True
                        else:
                            print("Ejeren har ingen adressebeskyttelse")
                            BoolBeskyttet = False
                    except ValueError:
                        print("Invalid data format for Beskyttelse")

                print(f"CVR: {CVR}")
                print(f"EjerNavn: {EjerNavn}")
                print(f"EjerAdresse: {EjerAdresse}")
                # Break loop if Hovedejer found
                break

        else:
            if "Person" in str(row["EjerType"]):
                print("Ejer er en person")
                if "Hovedejer" in str(row["EjerStatusKode_T"]):  # Check if Hovedejer
                    CPR = str(row["EjendePersonPersonNR"])
                    EjerAdresse = str(row["EjersAdresseDanmark"])
                    EjerNavn = str(row["NavnJusteret"])
                    BoolCPR = True

                    # Check for Beskyttelse (NULL check)
                    if pd.isnull(row["Beskyttelse"]):
                        print("Ejeren har ingen adressebeskyttelse")
                        BoolBeskyttet = False
                    else:
                        try:
                            if int(row["Beskyttelse"]) == 1:
                                print("Ejeren har adressebeskyttelse")
                                BoolBeskyttet = True
                            else:
                                print("Ejeren har ingen adressebeskyttelse")
                                BoolBeskyttet = False
                        except ValueError:
                            print("Invalid data format for Beskyttelse")

                    print(f"EjerNavn: {EjerNavn}")
                    print(f"EjerAdresse: {EjerAdresse}")
                    # Break loop if Hovedejer found
                    break

    
    if BoolCVR: 
        if AarhusKommuneCVR in CVR: 
            IsCvrAarhusKommune = True
            print("Aarhus Kommune er ejer på sagen - Sender fejlmail til Byggeri")
            #Sender error email til byggeri: 
            sender = "Rykkerbob<rpamtm001@aarhus.dk>" # Replace with actual sender
            subject = f"Sagsnummer: {Sagsnummer} har Aarhus Kommune som ejer"
            
            body = f"""Kære sagsbehandler, <br><br>
            Følgende sagsnummer: {Sagsnummer} har Aarhus Kommune som ejer, og mangler at få udsendt: {RykkerNummer}. rykker via Digital Post til ejeren.<br><br>
            Den nye frist dato er: {Dato}<br><br>
            I bedes derfor manuelt udsende rykkeren til ejeren via Digital Post<br><br>
            Med venlig hilsen<br><br>
            Teknik og Miljø<br><br>
            Digitalisering<br><br>
            Aarhus Kommune.
            """
            
            smtp_server = "smtp.adm.aarhuskommune.dk"   # Replace with your SMTP server
            smtp_port = 25                    # Replace with your SMTP port

            # Call the send_email function
            send_email(
                receiver="Gujc@aarhus.dk",
                sender=sender,
                subject=subject,
                body=body,
                smtp_server=smtp_server,
                smtp_port=smtp_port,
                html_body=True
            )

            try:
                TransactionID = str(uuid.uuid4())
                Uuid= str(uuid.uuid4())
                Aktivitetsnavn = "Nyt materiale"
    
                StartDato = datetime.now().strftime("%Y-%m-%dT%H:%M:%S+00:00")

                url = f"{KMDNovaURL}/Task/Import?api-version=2.0-Case"

                payload = {
                    "common": {
                        "transactionId": TransactionID,
                        "uuid": Uuid
                    },
                    "caseUuid": CaseUuid,
                    "title": Aktivitetsnavn,
                    "description": f"{BeskrivelseTilEjer} - Aarhus Kommune er ejer",
                        "caseworker": {
                            "kspIdentity": {
                                "racfId": racfId,
                                "fullName": fullName,
                            }
                        },
                    "startDate": StartDato,
                    "TaskTypeName": "Aktivitet",
                    "statusCode": "S"
                }
                headers = {
                    "Authorization": f"Bearer {Token}",
                    "Content-Type": "application/json"
                }

                response = requests.post(url, json=payload, headers=headers)
                print(f"API Response: {response.status_code}")
                print(f"API Response: {response.text}")

            except Exception as api_error:
                print(f"Error occurred during API call: {api_error}")
    
    if not EjerAdresse or EjerAdresse.strip() == "":
        print("Der findes ingen adresse på ejeren. Derfor bliver der ikke udsendt digital post")

        try:
            TransactionID = str(uuid.uuid4())
            Uuid= str(uuid.uuid4())
            Aktivitetsnavn = "Nyt materiale"

            StartDato = datetime.now().strftime("%Y-%m-%dT%H:%M:%S+00:00")

            url = f"{KMDNovaURL}/Task/Import?api-version=2.0-Case"

            payload = {
                "common": {
                    "transactionId": TransactionID,
                    "uuid": Uuid
                },
                "caseUuid": CaseUuid,
                "title": Aktivitetsnavn,
                "description": f"{BeskrivelseTilEjer} - Ejer mangler adresse",
                "caseworker": {
                    "kspIdentity": {
                        "racfId": racfId,
                        "fullName": fullName,
                    }
                },
                "startDate": StartDato,
                "TaskTypeName": "Aktivitet",
                "statusCode": "S"
            }
            headers = {
                "Authorization": f"Bearer {Token}",
                "Content-Type": "application/json"
            }

            response = requests.post(url, json=payload, headers=headers)
            print(f"API Response: {response.status_code}")

        except Exception as api_error:
            print(f"Error occurred during API call: {api_error}")


    else:
        if IsCvrAarhusKommune == False: 
            try: 
                if RykkerNummer == 2: 
                    print("Sender 2. rykker")
                    print(NewDate)
                    update_word_template(
                        EjerNavn,
                        EjerAdresse,
                        CaseTitle,
                        CaseAddress,
                        NewDate,
                        Sagsnummer,
                        DecisionDate = "",
                        input_filename="1. Orientering til ejer vedr. rykker for paabegyndelse.docx"
                    )
                    downloads_path = os.path.join("C:\\Users", os.getlogin(), "Downloads")
                    file_path = os.path.join(downloads_path, "1. Orientering til ejer vedr. rykker for paabegyndelse.docx")
                else: 
                    if RykkerNummer == 3: 
                        print("Sender 3. rykker")
                        update_word_template(
                            EjerNavn,
                            EjerAdresse,
                            CaseTitle,
                            CaseAddress,
                            NewDate,
                            Sagsnummer,
                            DecisionDate,
                            input_filename="2. Orientering til ejer vedr. rykker for paabegyndelse.docx"
                        )
                        downloads_path = os.path.join("C:\\Users", os.getlogin(), "Downloads")
                        file_path = os.path.join(downloads_path, "2. Orientering til ejer vedr. rykker for paabegyndelse.docx")
                    else: 
                        raise ValueError("Det er hverken 2. eller 3. rykker som er blevet oprettet")
                # Oploader dokument til Nova
                parsed_date = datetime.strptime(NewDate, "%d-%m-%Y")
                current_time = datetime.now().strftime("%H:%M:%S")
                ConfigDate = parsed_date.strftime("%Y-%m-%d") + "T" + current_time + "Z"
                novaUserId = "a1ac3c67-ab17-4969-8360-4989d0c8426a"
                fullName = "Robot SvcRpaMTM001"
                racfId = "AZMTM01"
                TransactionID = str(uuid.uuid4())
                DocumentID = str(uuid.uuid4())
                ID = str(uuid.uuid4())
                if RykkerNummer == 2: 
                    Title = "Orientering til ejer vedr. rykker for påbegyndelse"
                else:
                    if RykkerNummer == 3: 
                        Title = "2. Orientering til ejer vedr. rykker for påbegyndelse"
                    
                try: # Poster dokument

                    url = f"{KMDNovaURL}/Document/UploadFile/{TransactionID}/{DocumentID}?api-version=2.0-Case"

                    headers = {
                        "Authorization": f"Bearer {Token}",
                    }
                        # Open the file you want to upload
                    with open(file_path, "rb") as f:
                        files = {
                            "file": (
                                os.path.basename(file_path), 
                                f, 
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        }

                        response = requests.post(url, headers=headers, files=files)
               
                    print(f"API status: {response.status_code}")
                    print(f"API Response: {response.text}")


                except Exception as api_error:
                    print(f"Error occurred during API call: {api_error}")
                
                #Uoloader dokument: 
                try:

                    url = f"{KMDNovaURL}/Document/Import?api-version=2.0-Case"
                    payload = {
                        "common": {
                            "transactionId": ID,
                            "uuid": DocumentID
                        },
                        "caseUuid": CaseUuid,
                        "documentType": "Udgående",
                        "title": Title,
                        "approved": True,
                        "acceptReceived": True,
                        "AccessToDocuments": True,
                        "sensitivity": "Følsomme",
                        "documentDate": ConfigDate,
                        "caseworker": {
                            "kspIdentity": {
                                "novaUserId": novaUserId,
                                "racfId": racfId,
                                "fullName": fullName
                            }
                        },
                        "documentParty": {
                            "identificationType": identificationType,
                            "identification": identification
                        }
                    }
                    headers = {
                        "Authorization": f"Bearer {Token}",
                        "Content-Type": "application/json"
                    }

                    response = requests.post(url, json=payload, headers=headers)
                    print(f"API status: {response.status_code}")

                except Exception as api_error:
                    print(f"Error occurred during API call: {api_error}")

            except: 
                print("Sletter lokal fil")
                os.remove(file_path)
                raise Exception("Dokumentet blev ikke uploaded korrekt - sletter lokal fil")

            try: # Tjekker om documentet er uploaded 
                try:
                    url = f"{KMDNovaURL}/Document/GetList?api-version=2.0-Case"
                    payload = {
                            "common": {
                                "transactionId": str(uuid.uuid4())
                            },
                            "paging": {
                                "startRow": 1,
                                "numberOfRows": 1
                            },
                            "caseNumber": Sagsnummer,
                            "getOutput": {
                                "caseworker": True,
                                "documentDate": True,
                                "title": True,
                                "description": True,
                                "fileExtension": True,
                                "approved": True,
                                "acceptReceived": True
                            }
                    }                
                    headers = {
                            "Authorization": f"Bearer {Token}",
                            "Content-Type": "application/json"
                        }

                    response = requests.put(url, json=payload, headers=headers)
                    print(f"API status: {response.status_code}")
                except requests.exceptions.RequestException as e:
                    raise Exception(f"API request failed: {e}")
                
                Godkendt = "True"
                CurrentDate = parsed_date.strftime("%Y-%m-%d") + "T" + "00:00:00"
                # Extract the first document
                data = json.loads(response.text)  # Convert JSON string to dict
                item = data["documents"][0]

                # Check documentDate
                document_date_str = item.get("documentDate")
                if document_date_str is None:
                    DateMatches = False
                else:
                    print(f"Current date er: {CurrentDate}")
                    print(f"DocumentDate: {document_date_str}")
                    DateMatches = (document_date_str == CurrentDate)

                # Check title
                title_str = item.get("title")
                if title_str is None:
                    TitleMatches = False
                else:
                    TitleMatches = (title_str == Title)

                # Check approved
                approved_value = str(item.get("approved"))
                if approved_value is None:
                    ApprovedMatches = False
                else:
                    print(approved_value)
                    ApprovedMatches = (approved_value == Godkendt)

                # Output results
                print("DateMatches:", DateMatches)
                print("TitleMatches:", TitleMatches)
                print("ApprovedMatches:", ApprovedMatches)
                
                if ApprovedMatches and DateMatches and TitleMatches:
                    print("Dokumentet er sendt til modtageren, alt er godt")
                    DocumentSendt = True
                else:
                    DocumentSendt = False
            except: 
                os.remove(file_path)
                raise Exception("Dokumentet blev ikke uploaded korrekt - sletter lokal fil")
            try:
                if DocumentSendt: 
                    try:
                        # Setup ChromeDriver options
                        print("Initializing Chrome Driver...")
                        app_data_path = os.getenv("LOCALAPPDATA")
                        chrome_user_data_path = os.path.join(app_data_path, "Google", "Chrome", "User Data")

                        options = Options()
                        #options.add_argument("--headless=new")
                        options.add_argument(f"--user-data-dir={chrome_user_data_path}")
                        options.add_argument("--window-size=1920,900")
                        options.add_argument("--start-maximized")
                        options.add_argument("force-device-scale-factor=0.5")
                        options.add_argument("--disable-extensions")
                        options.add_argument("--profile-directory=Default")
                        options.add_argument("--remote-debugging-port=9222")
                        options.add_argument('--remote-debugging-pipe')

                        driver = webdriver.Chrome(options=options)
                        wait = WebDriverWait(driver, 120)

                        print("Creating WebDriver...")
                        print("Navigating to URL...")
                        driver.get("https://cap-wsswlbs-wm3q2021.kmd.dk/KMD.YH.KMDLogonWEB.NET/AspSson.aspx?KmdLogon_sApplCallback=https://cap-awswlbs-wm3q2021.kmd.dk/KMDNovaESDH/forside&KMDLogon_sProtocol=tcpip&KMDLogon_sApplPrefix=--&KMDLogon_sOrigin=ApplAsp&ExtraData=true")
                        print("Maximizing window...")
                        driver.maximize_window()

                        # Log in
                        print("Waiting for Username field...")
                        wait.until(EC.visibility_of_element_located((By.NAME, "UserInfo.Username"))).send_keys(RobotUserName)
                        print("Entered Username.")
                        wait.until(EC.visibility_of_element_located((By.NAME, "UserInfo.Password"))).send_keys(RobotPassword)
                        print("Entered Password.")
                        wait.until(EC.element_to_be_clickable((By.ID, "logonBtn"))).click()
                        print("Clicked Logon button.")

                        # Navigate
                        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button.btn.dropdown-toggle.object-select"))).click()
                        print("Opened dropdown.")
                        wait.until(EC.element_to_be_clickable((By.XPATH, "//li/a[.//span[text()='Sag']]"))).click()
                        print("Selected 'Sag'.")
                        wait.until(EC.visibility_of_element_located((By.ID, "SearchObject"))).send_keys(Sagsnummer + Keys.ENTER)
                        print("Entered Case Number.")
                        
                        text_to_find = "Orientering til ejer vedr. rykker for påbegyndelse"
                        wait.until(EC.element_to_be_clickable((By.XPATH, f"//span[contains(text(), '{text_to_find}')]"))).click()
                        print("Found and clicked target text.")

                        # Compare dates
                        input_date_string = NewDate
                        date_matches = False

                        time.sleep(8)
                        try:
                            input_date = datetime.strptime(input_date_string, "%d-%m-%Y").date()

                            date_element = driver.find_element(By.XPATH, "//label[text()='Dato']/following-sibling::div[@class='no-wrap ng-binding']")
                            found_date_string = date_element.text.strip()
                            found_date = datetime.strptime(found_date_string, "%d-%m-%Y").date()

                            if input_date == found_date:
                                print("The dates match.")
                                date_matches = True
                        except Exception as ex:
                            print("Error parsing date:", str(ex))
                            date_matches = False

                        if date_matches:
                            try:
                                print("Clicking 'document_details_show_multi_function'...")
                                wait.until(EC.element_to_be_clickable((By.ID, "document_details_show_multi_function"))).click()

                                print("Clicking send digital post...")
                                wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "span.test-send-digital-post"))).click()

                                time.sleep(4)

                                print("Clicking select all recipients...")
                                driver.find_element(By.XPATH, "//nova-checkbox[@ng-model='$ctrl.allSelectedRecipients' and @ng-change='$ctrl.toggleSelectAll()']//input").click()

                                print("Clicking remove recipients...")
                                wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@ng-click='$ctrl.removeSelectedRecipients()' and @class='btn' and @uib-tooltip='Fjern modtager']"))).click()

                                print("Clicking object select...")
                                wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@ng-click='$ctrl.objectTypeaheadSvc.clearIfTouched()' and @class= 'btn dropdown-toggle object-select']"))).click()

                                actions = ActionChains(driver)

                                if BoolCPR:
                                    print("CPR-nummer anvendes")
                                    actions.send_keys(Keys.DOWN).send_keys(Keys.ENTER).perform()
                                    wait.until(EC.visibility_of_element_located((By.ID, "viewModel_NewRecipient"))).send_keys(CPR)
                                    time.sleep(3)
                                    actions.send_keys(Keys.ENTER).perform()
                                    #wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button.btn-dialog-primary.test-send"))).click()
                                elif BoolCVR:
                                    print("CVR-nummer anvendes")
                                    actions.send_keys(Keys.DOWN).send_keys(Keys.DOWN).send_keys(Keys.ENTER).perform()
                                    wait.until(EC.visibility_of_element_located((By.ID, "viewModel_NewRecipient"))).send_keys(CVR)
                                    time.sleep(3)
                                    actions.send_keys(Keys.ENTER).perform()
                                    #wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button.btn-dialog-primary.test-send"))).click()
                                else:
                                    raise Exception("Business Rule Exception: Det er hverken et CPR eller CVR nummer")

                            except Exception as e:
                                print("Error during recipient handling:", str(e))
                        else:
                            raise Exception("Business Rule Exception: Date does not match.")

                        # Clean up
                        driver.quit()

                    except Exception as ex:
                        print("Exception occurred:", str(ex))
                        traceback.print_exc()
                        driver.quit()
                        raise     

                if BoolBeskyttet and DocumentSendt: 
                    print("Opdaterer dokumenttitlen med FORTROLIG")
                    FortroligTitle = f"FORTROLIG - {Title}"
                    try:
                        url = f"{KMDNovaURL}/Document/Update?api-version=2.0-Case"
                        payload = {
                            "common": {
                                "transactionId": TransactionID,
                                "uuid": DocumentID
                            },
                            "title": FortroligTitle
                        }             
                        headers = {
                                "Authorization": f"Bearer {Token}",
                                "Content-Type": "application/json"
                            }

                        response = requests.patch(url, json=payload, headers=headers)
                        print(f"API status: {response.status_code}")
                        print(f"API Response: {response.text}")
                    except requests.exceptions.RequestException as e:
                        raise Exception(f"API request failed: {e}")
                    
                    # Define email details
                    sender = "RykkerBob<rpamtm001@aarhus.dk>" 
                    subject = f"Sagsnummer: {Sagsnummer} har adressebeskyttelse"
                    body = f"""Kære sagsbehandler,<br><br>
                    Følgende sagsnummer: {Sagsnummer} har adressebeskyttelse - tjek op på dette. <br><br>
                    Med venlig hilsen<br><br>
                    Teknik & Miljø<br><br>
                    Digitalisering<br><br>
                    Aarhus Kommune
                    """
                    smtp_server = "smtp.adm.aarhuskommune.dk"   
                    smtp_port = 25               

                    # Call the send_email function
                    send_email(
                        receiver="Gujc@aarhus.dk",
                        sender=sender,
                        subject=subject,
                        body=body,
                        smtp_server=smtp_server,
                        smtp_port=smtp_port,
                        html_body=True
                    )
            
            except: 
                os.remove(file_path)
                raise Exception("Dokumentet blev ikke uploaded korrekt - sletter lokal fil")  
                  
            print(f"Sletter: {file_path}")
            os.remove(file_path)

    return {
        "out_Dokumentsendt": DocumentSendt,
    }